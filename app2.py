from flask import Flask, render_template, request,session,redirect,flash,url_for
from flask_sqlalchemy import SQLAlchemy
from flask_mail import Mail
from datetime import datetime
import json
import os
import pandas as pd
import os, sys, shutil, time
import joblib
import urllib.request
from geopy.geocoders import Nominatim

# import count_vect

from flask import Flask, jsonify, request
import numpy as np

import pandas as pd
import numpy as np
###################################
from flask import Flask, request, render_template, jsonify
import json
import datetime
import hashlib
from flask_mysqldb import MySQL
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.decomposition import PCA
from sklearn.ensemble import IsolationForest

import pandas as pd  
import numpy as np

import re
from flask import Flask, render_template, request
import os  # Add this import
import numpy as np

with open('config.json', 'r') as c:
    params = json.load(c)["params"]

local_server = True
app = Flask(__name__,template_folder='templates')
app.secret_key = 'super-secret-key'

app.config['MAIL_SERVER'] = 'smtp.googlemail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = params['gmail_user']
app.config['MAIL_PASSWORD'] = params['gmail_password']
mail = Mail(app)



if(local_server):
    app.config['SQLALCHEMY_DATABASE_URI'] = params['local_uri']
else:
    app.config['SQLALCHEMY_DATABASE_URI'] = params['prod_uri']

db = SQLAlchemy(app)

class Register(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=False)
    uname = db.Column(db.String(50), nullable=False)
    mobile = db.Column(db.String(10), nullable=False)
    email = db.Column(db.String(50), nullable=False)
    password = db.Column(db.String(10), nullable=False)
    cpassword = db.Column(db.String(10), nullable=False)

class Contact(db.Model):
    id=db.Column(db.Integer,primary_key=True)
    name=db.Column(db.String(50),nullable=False)
    email=db.Column(db.String(50),nullable=False)
    subject=db.Column(db.String(50),nullable=False)
    message=db.Column(db.String(250),nullable=False)

@app.route("/")
def Home():
    return render_template('index1.html',params=params)


@app.route("/about")
def About():
    return render_template('about.html',params=params)

@app.route("/contact",  methods=['GET','POST'])
def contact():
    if(request.method =='POST'):
        name=request.form.get('name')
        email=request.form.get('email')
        subject=request.form.get('subject')
        message=request.form.get('message')
        entry=Contact(name=name,email=email,subject=subject,message=message)
        db.session.add(entry)
        db.session.commit()
    return render_template('contact.html',params=params)

@app.route("/register",  methods=['GET','POST'])
def register():
    if(request.method=='POST'):
        name = request.form.get('name')
        uname = request.form.get('uname')
        mobile = request.form.get('mobile')
        email= request.form.get('email')
        password= request.form.get('password')
        cpassword= request.form.get('cpassword')

        user=Register.query.filter_by(email=email).first()
        if user:
            flash('Account already exist!Please login','success')
            return redirect(url_for('register'))
        if not(len(name)) >3:
            flash('length of name is invalid','error')
            return redirect(url_for('register')) 
        if (len(mobile))<10:
            flash('invalid mobile number','error')
            return redirect(url_for('register')) 
        if (len(password))<8:
            flash('length of password should be greater than 7','error')
            return redirect(url_for('register'))
        else:
             flash('You have registtered succesfully','success')
            
        entry = Register(name=name,uname=uname,mobile=mobile,email=email,password=password,cpassword=cpassword)
        db.session.add(entry)
        db.session.commit()
    return render_template('register.html',params=params)

@app.route("/login",methods=['GET','POST'])
def login():
    if (request.method== "GET"):
        if('email' in session and session['email']):
            return render_template('dashboard.html',params=params)
        else:
            return render_template("login.html", params=params)

    if (request.method== "POST"):
        email = request.form["email"]
        password = request.form["password"]
        
        login = Register.query.filter_by(email=email, password=password).first()
        if login is not None:
            session['email']=email
            return render_template('dashboard.html',params=params)
        else:
            flash("plz enter right password",'error')
            return render_template('login.html',params=params)


@app.route("/logout", methods = ['GET','POST'])
def logout():
    session.pop('email')
    return redirect(url_for('Home')) 



@app.route('/table', methods=['GET'])
def display_table():
    # Load the predictions from the CSV file
    df = pd.read_csv('predictions.csv')

    # Create a dictionary to store questions under each label
    questions_by_label = {label: [] for label in df['Predicted_Label'].unique()}

    # Populate the dictionary with questions
    for index, row in df.iterrows():
        questions_by_label[row['Predicted_Label']].append(row['Text'])

    return render_template('table.html', questions_by_label=questions_by_label)

from flask import Flask, render_template, request, send_file
import joblib
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random  # Add this line
# Load your trained model
model = joblib.load('your_model_svm.joblib')

# Assuming you have a preprocess_text function
def preprocess_text(text):
    # Your text preprocessing logic here
    return text

def add_numbered_paragraph(doc, text, font_size=12):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(font_size)
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return paragraph

@app.route('/predict', methods=['POST'])
def predict():
    if request.method == 'POST':
        # Get the input text from the form
        text = request.form['text']
        # Preprocess the input text
        processed_text = preprocess_text(text)
        # Make a prediction
        prediction = model.predict([processed_text])[0]
        return render_template('dashboard.html', text=text, prediction=prediction)

@app.route('/batch_predict', methods=['POST'])
def batch_predict():
    if request.method == 'POST':
        # Get the CSV file from the request
        csv_file = request.files['csv_file']
        # Load the CSV file into a DataFrame
        df = pd.read_csv(csv_file)
        # Preprocess the text in the DataFrame
        df['Processed_Text'] = df['Text'].apply(preprocess_text)
        # Make predictions for the entire DataFrame
        df['Predicted_Label'] = model.predict(df['Processed_Text'])
        # Create a new CSV file with predictions
        df.to_csv('predictions.csv', index=False)
        return 'Predictions saved to predictions.csv'



"""
@app.route('/generate_paper', methods=['GET'])
def generate_paper():
    # Load the predictions from the CSV file
    df = pd.read_csv('predictions.csv')

    # Create a dictionary to store questions under each label
    questions_by_label = {label: [] for label in df['Predicted_Label'].unique()}

    # Populate the dictionary with questions
    for index, row in df.iterrows():
        questions_by_label[row['Predicted_Label']].append(row['Text'])

    # Select a limited number of questions from each category
    selected_questions = {}
    max_questions_per_category = 4

    for label, questions in questions_by_label.items():
        selected_questions[label] = random.sample(questions, min(len(questions), max_questions_per_category))

    # Create a Word document for the question paper
    document = Document()

    # Add a title to the document
    title = document.add_heading('Question Paper', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add sequential numbered questions
    question_number = 1
    for label, questions in selected_questions.items():
        for question in questions:
            add_numbered_paragraph(document, f'{question_number}. {question} ({label})', font_size=12)
            question_number += 1

    # Save the document
    document.save('question_paper.docx')

    return 'Question paper generated!'
    """












"""
    
from docx.shared import Inches

@app.route('/generate_paper', methods=['GET'])
def generate_paper():
    # Load the predictions from the CSV file
    df = pd.read_csv('predictions.csv')

    # Create a dictionary to store questions under each label
    questions_by_label = {label: [] for label in df['Predicted_Label'].unique()}

    # Populate the dictionary with questions
    for index, row in df.iterrows():
        questions_by_label[row['Predicted_Label']].append(row['Text'])

    # Select a limited number of questions from each category
    selected_questions = {}
    max_questions_per_category = 4

    for label, questions in questions_by_label.items():
        selected_questions[label] = random.sample(questions, min(len(questions), max_questions_per_category))

    # Create a Word document for the question paper
    document = Document()

    # Add a title to the document
    title = document.add_heading('Question Paper', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add sequential numbered questions
    question_number = 1
    for label, questions in selected_questions.items():
        category_heading = document.add_heading(level=1)
        category_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        category_heading.add_run(f'{label} Questions\n').bold = True
        for question in questions:
            document.add_paragraph(f'{question_number}. {question}')
            question_number += 1

    # Change document format
    document.sections[0].page_height = Inches(11)
    document.sections[0].page_width = Inches(8.5)
    document.sections[0].left_margin = Inches(1)
    document.sections[0].right_margin = Inches(1)
    document.sections[0].top_margin = Inches(1)
    document.sections[0].bottom_margin = Inches(1)

    # Save the document
    document.save('question_paper.docx')

    return 'Question paper generated!'


"""








"""
from docx.shared import Pt

@app.route('/generate_paper', methods=['GET'])
def generate_paper():
    # Load the predictions from the CSV file
    df = pd.read_csv('predictions.csv')

    # Create a dictionary to store questions under each label
    questions_by_label = {
        'Knowledge': [],
        'Comprehension': [],
        'Application': [],
        'Analysis': [],
        'Synthesis': [],
        'Evaluation': []
    }

    # Populate the dictionary with questions
    for index, row in df.iterrows():
        label = row['Predicted_Label']
        if label in questions_by_label:
            questions_by_label[label].append(row['Text'])

    # Create a Word document for the question paper
    document = Document()

    # Add a title to the document
    title = document.add_heading('Question Paper', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add questions for each category
    question_number = 1
    for category, questions in questions_by_label.items():
        if questions:
            # Add category marker
            document.add_heading(f'{category} Questions', level=2)
            # Add questions
            for question in questions:
                document.add_paragraph(f'{question_number}. {question}', style='BodyText')
                question_number += 1

    # Add mark and mention the number of questions for each category
    mark_text = 'Mark:\n'
    for category, questions in questions_by_label.items():
        mark_text += f'{category}: {len(questions)}\n'
    document.add_paragraph(mark_text)

    # Change font size of the mark section
    for paragraph in document.paragraphs[-1:]:
        for run in paragraph.runs:
            run.font.size = Pt(10)

    # Save the document
    document.save('question_paper.docx')

    return 'Question paper generated!'

"""








from docx.shared import Inches

@app.route('/generate_paper', methods=['GET'])
def generate_paper():
    # Load the predictions from the CSV file
    df = pd.read_csv('predictions.csv')

    # Define categories and their respective marks
    categories = {
        'Knowledge': 1,
        'Comprehension': 2,
        'Application': 3,
        'Analysis': 4,
        'Synthesis': 5,
        'Evaluation': 6
    }

    # Create a dictionary to store questions under each label
    questions_by_label = {label: [] for label in categories}

    # Populate the dictionary with questions
    for index, row in df.iterrows():
        if row['Predicted_Label'] in categories:
            questions_by_label[row['Predicted_Label']].append(row['Text'])

    # Create a Word document for the question paper
    document = Document()

    # Add a title to the document
    title = document.add_heading('Question Paper', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add questions by category
    for category, mark in categories.items():
        if questions_by_label[category]:
            category_heading = document.add_heading(level=1)
            category_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            category_heading.add_run(f'{category} Questions (Marks: {mark})\n').bold = True
            for question_number, question in enumerate(questions_by_label[category], start=1):
                document.add_paragraph(f'{question_number}. {question}')

    # Add mention of marks and last category
    document.add_paragraph(f'\n\nNote: Marks mentioned in parentheses after each category.')
    document.add_paragraph(f'Questions for the last category are mentioned here.')

    # Change document format
    document.sections[0].page_height = Inches(11)
    document.sections[0].page_width = Inches(8.5)
    document.sections[0].left_margin = Inches(1)
    document.sections[0].right_margin = Inches(1)
    document.sections[0].top_margin = Inches(1)
    document.sections[0].bottom_margin = Inches(1)

    # Save the document
    document.save('question_paper.docx')

    return 'Question paper generated!'










"""
from docx.shared import Inches

@app.route('/generate_paper', methods=['GET'])
def generate_paper():
    # Load the predictions from the CSV file
    df = pd.read_csv('predictions.csv')

    # Define categories and their respective marks
    categories = {
        'Knowledge': 1,
        'Comprehension': 2,
        'Application': 3,
        'Analysis': 4,
        'Synthesis': 5,
        'Evaluation': 6
    }

    # Create a dictionary to store questions under each label
    questions_by_label = {label: [] for label in categories}

    # Populate the dictionary with questions
    for index, row in df.iterrows():
        if row['Predicted_Label'] in categories:
            questions_by_label[row['Predicted_Label']].append(row['Text'])

    # Create a Word document for the question paper
    document = Document()

    # Add a title to the document
    title = document.add_heading('Question Paper', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add questions by category
    for category, mark in categories.items():
        if questions_by_label[category]:
            for question_number, question in enumerate(questions_by_label[category], start=1):
                document.add_paragraph(f'{question_number}. {question} (Marks: {mark}, Category: {category})')

    # Add mention of marks at the end
    document.add_paragraph(f'\n\nNote: Marks mentioned in parentheses at the end of each question.')

    # Change document format
    document.sections[0].page_height = Inches(11)
    document.sections[0].page_width = Inches(8.5)
    document.sections[0].left_margin = Inches(1)
    document.sections[0].right_margin = Inches(1)
    document.sections[0].top_margin = Inches(1)
    document.sections[0].bottom_margin = Inches(1)

    # Save the document
    document.save('question_paper.docx')

    return 'Question paper generated!'
"""










@app.route('/download_paper', methods=['GET'])
def download_paper():
    return send_file('question_paper.docx', as_attachment=True)
#############
if __name__ == '__main__':
    app.run(debug=True)